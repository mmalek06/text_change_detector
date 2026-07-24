import threading
import zlib

import numpy as np
from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MINILM = "sentence-transformers/all-MiniLM-L6-v2"
RT_FOOTNOTES = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
CT_FOOTNOTES = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
FOOTNOTE_MARKER = "SIDENOTE_MARKER_UNIQUE_TOKEN"


def add_footnote(doc, paragraph, note_text, note_id=1):
    xml = (
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>"
        f'<w:footnote w:id="{note_id}"><w:p><w:r>'
        f'<w:t xml:space="preserve">{note_text}</w:t>'
        "</w:r></w:p></w:footnote>"
        "</w:footnotes>"
    ).encode("utf-8")
    part = Part(PackURI("/word/footnotes.xml"), CT_FOOTNOTES, blob=xml, package=doc.part.package)

    doc.part.relate_to(part, RT_FOOTNOTES)

    run = paragraph.add_run()
    ref = OxmlElement("w:footnoteReference")

    ref.set(qn("w:id"), str(note_id))
    run._r.append(ref)


def build_legal_docx(path):
    doc = Document()

    doc.add_heading("1. Lease Agreement", level=1)
    doc.add_heading("1.1 Rent", level=2)
    doc.add_paragraph("The tenant pays the monthly rent by the fifth day of each month.")
    doc.add_paragraph("Late payments incur a penalty of five percent of the outstanding balance.")
    doc.add_heading("2. Obligations", level=1)
    doc.add_paragraph("The following obligations apply to both parties:")
    doc.add_paragraph("the deposit is returned within thirty days of termination", style="List Number")

    table = doc.add_table(rows=2, cols=2)

    table.cell(0, 0).text = "Party"
    table.cell(0, 1).text = "Duty"
    table.cell(1, 0).text = "Tenant"
    table.cell(1, 1).text = "The tenant maintains the flat in good repair."
    doc.save(str(path))

    return path


def build_it_docx(path):
    doc = Document()

    doc.add_heading("1. Authentication", level=1)
    doc.add_paragraph("The service authenticates each request using a bearer token.")
    doc.add_paragraph("The gateway rejects requests that present an expired token.")
    doc.add_heading("2. Storage", level=1)
    doc.add_paragraph("The database replicates writes to two standby nodes.")

    form = doc.add_table(rows=2, cols=2)

    form.cell(0, 0).text = "Engine"
    form.cell(0, 1).text = "PostgreSQL"
    form.cell(1, 0).text = "Replication"
    form.cell(1, 1).text = "The cluster streams the write-ahead log to every replica."
    doc.save(str(path))

    return path


def build_toc_docx(path):
    doc = Document()

    doc.add_heading("Table of Contents", level=1)

    for line in ("1\tIntroduction\t1", "1.1\tScope\t2", "2\tArchitecture\t5"):
        doc.add_paragraph(line)

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("The system authenticates each request using a bearer token.")
    doc.save(str(path))

    return path


def build_table_docx(path):
    doc = Document()

    doc.add_heading("1. Configuration", level=1)

    headed = doc.add_table(rows=3, cols=2)

    headed.cell(0, 0).text = "Field"
    headed.cell(0, 1).text = "Description"
    headed.cell(1, 0).text = "timeout"
    headed.cell(1, 1).text = "The client aborts the request after thirty seconds."
    headed.cell(2, 0).text = "retries"
    headed.cell(2, 1).text = "The client retries a failed call three times."
    doc.add_heading("2. Runtime", level=1)

    form = doc.add_table(rows=3, cols=2)

    form.cell(0, 0).text = "Latency"
    form.cell(0, 1).text = "The proxy forwards each packet within two milliseconds."
    form.cell(1, 0).text = "Status"
    form.cell(1, 1).text = "Active"
    form.cell(2, 0).text = "Owner"
    form.cell(2, 1).text = "Finance"
    doc.save(str(path))

    return path


def build_footnote_docx(path):
    doc = Document()

    doc.add_heading("1. Definitions", level=1)

    para = doc.add_paragraph("The agreement binds both parties from the effective date.")

    add_footnote(doc, para, f"{FOOTNOTE_MARKER} the effective date is defined in schedule A.")
    doc.add_paragraph("The parties agree to resolve disputes through arbitration.")
    doc.save(str(path))

    return path


DISTANT_CHAPTERS = [
    ("1. Payment and Rent", [
        "The tenant pays the monthly rent by the fifth day of each month.",
        "Late payments incur a penalty of five percent of the outstanding balance.",
        "The landlord issues an invoice for every rental period.",
        "The tenant settles each invoice through a bank transfer.",
        "Unpaid rent accrues interest until the balance clears.",
        "A security deposit is collected before the tenancy begins.",
        "The deposit is refunded after the final inspection.",
        "Rent increases are announced sixty days in advance.",
        "The lease sets the payment schedule for the whole term.",
        "Partial payments are applied to the oldest outstanding balance first.",
        "A receipt is issued for every rent payment received.",
        "The tenant may pay rent by card, transfer, or standing order.",
        "Overdue rent is reported after a grace period of ten days.",
        "The landlord may withhold the deposit to cover unpaid rent.",
        "Rent is prorated for a partial month at move in.",
    ]),
    ("2. Authentication", [
        "The service authenticates each request using a bearer token.",
        "The gateway rejects requests that present an expired token.",
        "Administrators rotate the signing keys every ninety days.",
        "The login endpoint throttles repeated failed attempts.",
        "The system logs every authentication failure for review.",
        "Multi factor authentication is required for administrator accounts.",
        "Session tokens expire after thirty minutes of inactivity.",
        "A refresh token issues a new access token when valid.",
        "Revoked credentials are checked against a denylist.",
        "Passwords are hashed with a salted algorithm before storage.",
        "The identity provider issues signed assertions to clients.",
        "Single sign on federates login across internal tools.",
        "A brute force lockout blocks an account after five failures.",
        "Access scopes limit what each token is permitted to do.",
        "The audit trail records every login and logout event.",
    ]),
    ("3. Data Storage", [
        "The database replicates writes to two standby nodes.",
        "The engine compacts old segments during nightly maintenance.",
        "Each query reads from the nearest available replica.",
        "The cluster shards large tables across several disks.",
        "A background job backs up every partition to cold storage.",
        "Indexes are rebuilt when fragmentation crosses a threshold.",
        "Write ahead logs guarantee durability after a crash.",
        "Snapshots are retained for thirty days before deletion.",
        "Read replicas lag the primary by a few milliseconds.",
        "The storage tier moves cold rows to cheaper disks.",
        "Compression reduces the footprint of archived tables.",
        "A checksum verifies the integrity of every stored block.",
        "Failed nodes are rebuilt from the surviving replicas.",
        "The catalog tracks the location of every shard.",
        "Vacuuming reclaims space from deleted rows.",
    ]),
    ("4. Invoicing and Billing", [
        "Every invoice becomes due within fourteen days of issuance.",
        "The finance team charges a surcharge for overdue balances.",
        "Tenants settle their rent through the online billing portal.",
        "The portal emails a payment receipt after each transaction.",
        "Outstanding rent appears on the next monthly statement.",
        "A late fee is added to invoices past their due date.",
        "The billing cycle closes on the last day of the month.",
        "Refunds are issued to the original payment method.",
        "Each statement itemizes rent, fees, and adjustments.",
        "The portal accepts card payments and bank transfers.",
        "An invoice lists the amount due and the payment deadline.",
        "Unpaid invoices are escalated to a collections notice.",
        "Credit balances roll over to the following billing period.",
        "The system reconciles payments against open invoices nightly.",
        "A final statement is sent when the account is closed.",
    ]),
]


def build_distant_topics_docx(path):
    doc = Document()

    for title, sentences in DISTANT_CHAPTERS:
        doc.add_heading(title, level=1)

        for sentence in sentences:
            doc.add_paragraph(sentence)

    doc.save(str(path))

    return path


class StubEmbedder:
    def __init__(self, table=None, dim=8):
        self.table = dict(table or {})
        self.dim = dim
        self.calls = []

    def encode(self, sentences, normalize_embeddings=True):
        self.calls.append(list(sentences))
        rows = []

        for sentence in sentences:
            if sentence in self.table:
                vector = np.asarray(self.table[sentence], dtype=np.float64)
            else:
                vector = self._auto_vector(sentence)

            if normalize_embeddings:
                norm = np.linalg.norm(vector)

                if norm:
                    vector = vector / norm

            rows.append(vector)

        return np.asarray(rows, dtype=np.float64)

    def _auto_vector(self, sentence):
        vector = np.zeros(self.dim, dtype=np.float64)
        vector[zlib.crc32(sentence.encode("utf-8")) % self.dim] = 1.0

        return vector


class _BoundStub:
    def __init__(self, parent, schema):
        self.parent = parent
        self.schema = schema

    def invoke(self, prompt):
        self.parent.calls.append((self.schema, prompt))

        return self.parent.handlers[self.schema](prompt)


class StructuredLLMStub:
    """A stand-in for a LangChain chat model used with structured output.

    `handlers` maps each response schema to a callable `(prompt) -> BaseModel`.
    Every `with_structured_output(schema).invoke(prompt)` call is recorded on
    `.calls` as `(schema, prompt)` so tests can assert what was sent.
    """

    def __init__(self, handlers):
        self.handlers = dict(handlers)
        self.calls = []

    def with_structured_output(self, schema, **kwargs):
        return _BoundStub(self, schema)


class _BarrierBoundStub:
    def __init__(self, parent, schema):
        self.parent = parent
        self.schema = schema

    def invoke(self, prompt):
        self.parent.barrier.wait(timeout=self.parent.timeout)

        with self.parent.lock:
            self.parent.calls.append((self.schema, prompt))

        return self.parent.handlers[self.schema](prompt)


class BarrierLLMStub:
    """A structured-output stub whose every call waits at a shared barrier.

    With `parties` set to the expected concurrency, a call returns only once
    that many calls are in flight at the same moment, which proves genuine
    overlap. A caller that cannot overlap trips the barrier timeout instead,
    failing the test fast rather than hanging the suite.
    """

    def __init__(self, handlers, parties, timeout=2.0):
        self.handlers = dict(handlers)
        self.barrier = threading.Barrier(parties)
        self.timeout = timeout
        self.lock = threading.Lock()
        self.calls = []

    def with_structured_output(self, schema, **kwargs):
        return _BarrierBoundStub(self, schema)
