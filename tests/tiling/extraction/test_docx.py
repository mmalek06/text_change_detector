from docx import Document
from docx.shared import Pt

from text_change_detector.models import Segment
from text_change_detector.tiling.extraction.docx import (
    column_headers,
    extract_docx,
    has_numbering,
    heuristic_heading_level,
    is_toc_paragraph,
    median_font_size,
    outline_heading_level,
)
from text_change_detector.tiling.extraction.shared import is_content
from tests.helpers import FOOTNOTE_MARKER


def texts(segments):
    return [s.text for s in segments]


def payloads(segments):
    return [p for s in segments for p in s.payload]


class TestExtractLegal:
    def test_returns_segments(self, legal_docx, nlp):
        segments = extract_docx(legal_docx, nlp)
        assert segments
        assert all(isinstance(s, Segment) for s in segments)

    def test_every_segment_text_is_content(self, legal_docx, nlp):
        for segment in extract_docx(legal_docx, nlp):
            assert is_content(segment.text, nlp)

    def test_nested_heading_path_in_section(self, legal_docx, nlp):
        sections = {s.section for s in extract_docx(legal_docx, nlp)}

        assert "1. Lease Agreement > 1.1 Rent" in sections
        assert "2. Obligations" in sections

    def test_body_sentence_extracted(self, legal_docx, nlp):
        assert "The tenant pays the monthly rent by the fifth day of each month." in texts(extract_docx(legal_docx, nlp))

    def test_numbered_item_gets_lead_in_prefix(self, legal_docx, nlp):
        merged = [t for t in texts(extract_docx(legal_docx, nlp))
                  if t.startswith("The following obligations apply to both parties:")
                  and "deposit is returned" in t]

        assert len(merged) == 1

    def test_table_header_context_becomes_side_note(self, legal_docx, nlp):
        assert "Party: Tenant" in payloads(extract_docx(legal_docx, nlp))

    def test_table_header_labels_are_not_standalone_segments(self, legal_docx, nlp):
        assert "Party" not in texts(extract_docx(legal_docx, nlp))
        assert "Duty" not in texts(extract_docx(legal_docx, nlp))


class TestExtractIt:
    def test_sections(self, it_docx, nlp):
        sections = {s.section for s in extract_docx(it_docx, nlp)}

        assert "1. Authentication" in sections
        assert "2. Storage" in sections

    def test_content_sentences_extracted(self, it_docx, nlp):
        extracted = texts(extract_docx(it_docx, nlp))

        assert "The service authenticates each request using a bearer token." in extracted
        assert "The cluster streams the write-ahead log to every replica." in extracted


class TestExtractToc:
    def test_toc_lines_are_dropped(self, toc_docx, nlp):
        extracted = texts(extract_docx(toc_docx, nlp))

        assert all("\t" not in t for t in extracted)
        assert not any("Scope" in t or "Architecture" in t for t in extracted)

    def test_real_content_survives(self, toc_docx, nlp):
        segments = extract_docx(toc_docx, nlp)

        assert texts(segments) == ["The system authenticates each request using a bearer token."]
        assert segments[0].section == "1. Introduction"


class TestExtractTables:
    def test_header_table_content_and_context(self, table_docx, nlp):
        segments = extract_docx(table_docx, nlp)

        assert "The client aborts the request after thirty seconds." in texts(segments)
        assert "The client retries a failed call three times." in texts(segments)
        assert "Field: timeout" in payloads(segments)
        assert "Field: retries" in payloads(segments)

    def test_two_column_form_content_and_labels(self, table_docx, nlp):
        segments = extract_docx(table_docx, nlp)

        assert "The proxy forwards each packet within two milliseconds." in texts(segments)
        assert "Status: Active" in payloads(segments)
        assert "Owner: Finance" in payloads(segments)


class TestFootnotesCurrentlyIgnored:
    def test_footnote_text_is_not_extracted(self, footnote_docx, nlp):
        segments = extract_docx(footnote_docx, nlp)
        blob = " ".join(t + " " + " ".join(p for p in s.payload) for s, t in zip(segments, texts(segments)))

        assert FOOTNOTE_MARKER not in blob

    def test_referencing_body_text_survives(self, footnote_docx, nlp):
        extracted = texts(extract_docx(footnote_docx, nlp))

        assert "The agreement binds both parties from the effective date." in extracted
        assert "The parties agree to resolve disputes through arbitration." in extracted


class TestHeadingDetection:
    def test_outline_heading_level_from_style(self):
        doc = Document()

        assert outline_heading_level(doc.add_heading("Chapter", level=1)) == 1
        assert outline_heading_level(doc.add_heading("Section", level=2)) == 2

    def test_outline_level_none_for_body(self):
        doc = Document()

        assert outline_heading_level(doc.add_paragraph("A normal sentence of body text.")) is None

    def test_heuristic_detects_caps_bold_heading(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("SECURITY MODEL")
        run.bold = True

        assert heuristic_heading_level(para, body_font_size=11.0) == 1

    def test_heuristic_ignores_sentence_with_terminal_punctuation(self):
        doc = Document()
        para = doc.add_paragraph("The service authenticates each request using a bearer token.")

        assert heuristic_heading_level(para, body_font_size=11.0) is None

    def test_heuristic_numbered_depth(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("1.2 Access Control")
        run.bold = True

        assert heuristic_heading_level(para, body_font_size=11.0) == 2


class TestTocDetection:
    def test_tab_separated_toc_line(self):
        doc = Document()

        assert is_toc_paragraph(doc.add_paragraph("1.1\tScope\t2")) is True

    def test_normal_paragraph_is_not_toc(self):
        doc = Document()

        assert is_toc_paragraph(doc.add_paragraph("Just a normal sentence.")) is False


class TestNumbering:
    def test_list_number_style_has_numbering(self):
        doc = Document()

        assert has_numbering(doc.add_paragraph("an item", style="List Number")) is True

    def test_plain_paragraph_has_no_numbering(self):
        doc = Document()

        assert has_numbering(doc.add_paragraph("plain body")) is False


class TestColumnHeaders:
    def test_label_row_is_detected(self, nlp):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        table.cell(0, 0).text = "Field"
        table.cell(0, 1).text = "Description"
        table.cell(1, 0).text = "x"
        table.cell(1, 1).text = "The client aborts the request after thirty seconds."

        assert column_headers(table, nlp) == ["Field", "Description"]

    def test_content_first_row_is_not_headers(self, nlp):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        table.cell(0, 0).text = "Latency"
        table.cell(0, 1).text = "The proxy forwards each packet within two milliseconds."
        table.cell(1, 0).text = "Status"
        table.cell(1, 1).text = "Active"

        assert column_headers(table, nlp) is None

    def test_single_row_table_has_no_headers(self, nlp):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)

        table.cell(0, 0).text = "Field"
        table.cell(0, 1).text = "Value"

        assert column_headers(table, nlp) is None


class TestMedianFontSize:
    def test_median_of_run_sizes(self):
        doc = Document()

        for _ in range(3):
            doc.add_paragraph().add_run("x").font.size = Pt(11)

        doc.add_paragraph().add_run("y").font.size = Pt(20)

        assert median_font_size(doc) == 11.0

    def test_none_when_no_sizes(self):
        doc = Document()

        doc.add_paragraph("no explicit size")

        assert median_font_size(doc) is None
